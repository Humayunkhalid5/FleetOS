from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path
from datetime import date

OUT = Path(r"D:\FleetOs copy\FleetOS_System_Documentation.docx")
NAVY = "0D1B2A"
MID = "415A77"
BLUE = "173EAA"
PALE = "EAF0F7"
LIGHT = "F7F9FC"
MUTED = "5B6B7D"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    ind = OxmlElement('w:tblInd')
    ind.set(qn('w:w'), '120')
    ind.set(qn('w:type'), 'dxa')
    tblPr.append(ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcPr = cell._tc.get_or_add_tcPr()
            margins = tcPr.first_child_found_in('w:tcMar')
            if margins is None:
                margins = OxmlElement('w:tcMar')
                tcPr.append(margins)
            for side in ('top', 'start', 'bottom', 'end'):
                el = margins.find(qn(f'w:{side}'))
                if el is None:
                    el = OxmlElement(f'w:{side}')
                    margins.append(el)
                el.set(qn('w:w'), '90' if side in ('top', 'bottom') else '120')
                el.set(qn('w:type'), 'dxa')

def mark_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    trPr.append(header)

def set_run(run, size=11, color=NAVY, bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def set_para(p, before=0, after=6, line=1.1, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep

def add_text(doc, text, size=11, color=NAVY, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para(p, before, after)
    set_run(p.add_run(text), size, color, bold, italic)
    return p

def add_heading(doc, text, level=1):
    spec = {1:(16, BLUE, 16, 8), 2:(13, BLUE, 12, 6), 3:(11.5, MID, 8, 4)}[level]
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    set_para(p, spec[2], spec[3], 1.05, True)
    set_run(p.add_run(text), spec[0], spec[1], True)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        set_para(p, 0, 3, 1.15)
        set_run(p.add_run(item), 10.5, NAVY)

def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        set_para(p, 0, 4, 1.15)
        set_run(p.add_run(item), 10.5, NAVY)

def add_callout(doc, title, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [TABLE_WIDTH])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    set_para(p, 1, 2, 1.12)
    set_run(p.add_run(title + '  '), 10.5, BLUE, True)
    set_run(p.add_run(text), 10.5, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, PALE)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.0)
        set_run(p.add_run(head), 9.5, NAVY, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_para(p, 0, 0, 1.08)
            set_run(p.add_run(str(value)), 9.3, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table

def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.header_distance = Inches(0.45); section.footer_distance = Inches(0.42)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(NAVY)
    for name in ('Heading 1', 'Heading 2', 'Heading 3'):
        style = doc.styles[name]
        style.font.name = 'Calibri'; style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(header, 0, 0, 1.0)
    set_run(header.add_run('FLEETOS | SYSTEM DOCUMENTATION'), 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, 0, 0, 1.0)
    set_run(footer.add_run('FleetOS - Internal project documentation'), 8.5, MUTED)

def cover(doc):
    add_text(doc, 'FLEETOS', 13, BLUE, True, before=45, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'System Documentation & Test Guide', 28, NAVY, True, before=0, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'Architecture, portal workflow, permissions, MongoDB data handling, deployment, and acceptance testing', 13, MID, False, before=0, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc, 'Purpose', 'A complete handover guide for instructors, developers, testers, and project administrators. It describes how the three FleetOS portals work together and how to verify that every important operation persists in MongoDB.', 'EAF0F7')
    add_text(doc, 'Project scope', 12, NAVY, True, before=24, after=8)
    add_bullets(doc, [
        'Marketplace for approved Pakistani companies offering products and services.',
        'Client portal for discovery, booking, chat, tracking, payment selection, and reviews.',
        'Company portal for company profile, requests, technicians, inventory, services, customers, chat, analytics, and payments.',
        'Separate Super Admin portal for approval, company control, users, reviews, bookings, finance, support, security, and audit history.',
    ])
    add_text(doc, f'Prepared: {date.today().strftime("%d %B %Y")}', 10.5, MUTED, before=20, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)

def document_body(doc):
    add_heading(doc, '1. Platform Overview', 1)
    add_text(doc, 'FleetOS is a MongoDB-backed MERN marketplace and service-management system. It has three separate role experiences that exchange data through Express APIs and Socket.IO. MongoDB is the source of truth; browser memory is only used for the currently signed-in session and page state.', after=8)
    add_table(doc, ['Portal', 'Primary users', 'Main purpose', 'Local address'], [
        ['Client portal', 'Customers / clients', 'Discover approved companies, book, chat, track, pay, and review.', 'http://localhost:5173'],
        ['Company portal', 'Approved company owners and staff', 'Manage the company profile, offers, operations, and client work.', 'http://localhost:5173'],
        ['Super Admin portal', 'One protected Super Admin account', 'Approve, oversee, secure, and audit the whole platform.', 'http://localhost:5174'],
        ['API and real-time server', 'All portals', 'Authentication, data operations, business rules, and chat updates.', 'http://localhost:5000'],
    ], [1600, 1900, 4080, 1780])
    add_callout(doc, 'Pakistan catalogue', 'The development seed maintains a Pakistan-only city catalogue. Public company discovery shows only approved and active companies; pending, rejected, or suspended companies are not visible to clients.', 'F7F9FC')

    add_heading(doc, '2. Technology and Components', 1)
    add_table(doc, ['Layer', 'Technology / component', 'Responsibility'], [
        ['Client and company UI', 'React + Vite', 'Role-based pages, forms, responsive UI, API calls, and Socket.IO chat/tracking updates.'],
        ['Super Admin UI', 'Separate React + Vite application', 'Independent Admin login and protected governance console.'],
        ['Backend', 'Node.js + Express', 'REST API, validation, authorization, business workflow, and error handling.'],
        ['Database', 'MongoDB + Mongoose', 'Persistent records, references, indexes, validations, and data retrieval.'],
        ['Real time', 'Socket.IO', 'Conversation and work-status updates for authorized users.'],
        ['Security', 'bcryptjs + JWT cookies', 'Password hashing, signed sessions, session invalidation, and protected APIs.'],
    ], [1500, 2600, 5260])
    add_text(doc, 'Important configuration values', 12, NAVY, True, before=8, after=4)
    add_bullets(doc, [
        'MONGODB_URI: database connection string. Development normally uses mongodb://127.0.0.1:27017/fleetos.',
        'JWT_SECRET: required production secret used to sign and validate sessions.',
        'CORS_ORIGINS: required in production; contains the deployed client and Admin website origins.',
        'NODE_ENV=production: prevents development demonstration data from being inserted or overwritten.',
    ])

    add_heading(doc, '3. Data Model and MongoDB Ownership', 1)
    add_text(doc, 'The following records are stored in MongoDB and connected by references. Company-scoped controllers obtain the company from the signed-in user session, rather than trusting a company ID sent by the browser.', after=6)
    add_table(doc, ['Collection', 'Key data', 'Created / changed by'], [
        ['User', 'name, email, hashed password, role, status, company link, session version', 'Registration, profile edits, Admin account controls'],
        ['Company', 'business identity, city, areas, logo, cover image, description, approval status, rating', 'Company registration, Company Details, Admin approval/status'],
        ['Service', 'company, service ID, name, category, price, duration, active status', 'Approved company service management'],
        ['Technician', 'company, avatar, role, availability, ratings, experience', 'Approved company staff management'],
        ['Inventory', 'company, SKU, category, quantity, pricing, stock level', 'Approved company inventory management'],
        ['Booking', 'customer, company, service snapshot, technician, status, location, price, payment method', 'Client booking and company operations'],
        ['ChatMessage', 'booking / conversation, sender, text, timestamps', 'Client-company conversations'],
        ['Payment', 'booking, customer, company, amount, method, status', 'Checkout request or company payment recording'],
        ['Review', 'customer, company, booking, rating, comment, company reply', 'Client review and company response'],
        ['AuditEvent / SupportRequest / City / Customer', 'governance events, support work, Pakistan cities, company customer records', 'Admin, system seed, and company operations'],
    ], [1600, 4700, 3060])
    add_callout(doc, 'Persistence rule', 'If a user adds, edits, or deletes a company resource, the API writes the change to MongoDB first and then returns the stored result to the UI. A refresh or new login should display the latest saved record.', 'EAF0F7')

    add_heading(doc, '4. Roles and Permissions', 1)
    add_table(doc, ['Action', 'Client', 'Approved company', 'Pending / suspended company', 'Super Admin'], [
        ['Register and sign in', 'Yes', 'Yes after approval', 'Can sign in; operational APIs blocked', 'Separate protected login only'],
        ['Browse public companies', 'Yes', 'Public endpoint available', 'Public endpoint available', 'Yes'],
        ['Edit own profile', 'Yes', 'Company profile and branding', 'Basic account only', 'Admin security settings'],
        ['Create a booking', 'Yes', 'No', 'No', 'Oversight only'],
        ['Manage own services, staff, inventory, customers', 'No', 'Yes; company-scoped', 'No', 'View / control only'],
        ['Assign technician / update work status / record payment', 'No', 'Yes; own bookings only', 'No', 'View / control only'],
        ['Chat on authorized booking', 'Yes', 'Yes; own company conversations', 'No operations', 'Audit / support visibility'],
        ['Create review', 'Yes; eligible completed booking', 'No; can reply', 'No', 'Read all reviews'],
        ['Approve, reject, suspend, or reactivate companies', 'No', 'No', 'No', 'Yes; audit reason required'],
        ['View platform-wide users, payments, support, and audit history', 'No', 'Own data only', 'No operations', 'Yes'],
    ], [2580, 1450, 1820, 1920, 1590])
    add_text(doc, 'Authorization controls', 12, NAVY, True, before=8, after=4)
    add_bullets(doc, [
        'Every protected request validates the signed JWT session and checks that the account is active and its session version is current.',
        'Company operational routes additionally require role=company and approvalStatus=approved.',
        'Super Admin API routes use a different Admin session cookie and require role=super-admin.',
        'The Admin portal never accepts normal client or company accounts as Admin accounts.',
        'Admin mutations require a reason and create an audit event.',
    ])

    add_heading(doc, '5. End-to-End Business Flow', 1)
    add_numbered(doc, [
        'Company registration: a company submits business identity, city, areas, logo, business license, owner credentials, and description. The company record is created with pending approval.',
        'Admin review: Super Admin opens Registration Requests, checks the submitted record/document, supplies an audit reason, and approves or rejects it.',
        'Client discovery: an approved, active company becomes visible in the client marketplace. Its own logo, cover image, description, location, active services, and service categories are returned from MongoDB.',
        'Client booking: the client searches by city, area, text, or service category; opens the company, selects an active service and optional inventory items, chooses a location and payment method, and creates the booking.',
        'Company fulfillment: the selected company sees the request, accepts or updates it, assigns an available technician, updates tracking/status, communicates with the client, and records payment when appropriate.',
        'Client completion: the client follows status and tracking, sees the technician/company details, completes payment selection, and submits one review for an eligible completed booking.',
        'Platform oversight: Super Admin can view companies, users, reviews, jobs, payments, support, audit history, and can block a company or account where necessary.',
    ])
    add_heading(doc, '5.1 Detailed Operational Journey', 2)
    add_text(doc, 'The following is the complete expected workflow from a new company joining FleetOS to a completed client service. Every state-changing step is saved in MongoDB and becomes visible only to the correct authorized role.', after=6)
    add_table(doc, ['Stage', 'Actor and page action', 'System result / next user'], [
        ['1. Company applies', 'A business owner opens Company Registration, enters owner and business details, Pakistan city and service areas, adds a logo and license, writes a public description, and creates credentials.', 'A User and Company record are saved as pending. The company is not visible to clients and cannot use operational company APIs.'],
        ['2. Admin verifies', 'Super Admin signs into the separate Admin console, opens Registration Requests, reviews the details and license, then approves or rejects with a reason.', 'Approval status and audit event are saved. An approved active company can now access its portal and become client-visible.'],
        ['3. Company prepares its listing', 'The approved company opens Company Details, adds its own logo and client-card cover, description, exact location and areas. It then creates services, technicians with photos, and relevant inventory.', 'The client marketplace reads this exact company data. Cards, service tags, categories, public details, and images come from the company record and related services.'],
        ['4. Client finds a company', 'A client registers/signs in, then searches Browse Companies by company wording, city, area, or service category. The client opens a company card to read its actual description, offers, location, staff, and reviews.', 'Only approved, active companies matching the filters are returned. Pending, rejected, suspended, or unrelated companies are hidden.'],
        ['5. Client asks a question', 'From the company card/details or Messages page, the client starts a conversation and sends a message to the selected company.', 'The conversation is saved. The related company sees it in its Chat area; each party sees sender identity and conversation history after refresh.'],
        ['6. Client places a booking', 'The client selects one active company service, optional company inventory, location, preferred time/request notes, and cash or online payment choice. The client confirms the booking.', 'A Booking stores customer, company, service snapshot, price data, selected items, location, payment method, and initial Pending status. The chosen company receives the request.'],
        ['7. Company accepts and assigns', 'The company opens Requests/Bookings, reviews the client request, accepts it, and selects an available technician from its own Staff list.', 'The booking becomes Assigned and stores the technician reference. The client can now see technician identity and assignment status.'],
        ['8. Work is tracked', 'The company changes job status and tracking details as the technician travels and works: En Route, Arrived, In Progress, and Completed. Both sides can continue authorized chat.', 'The client Live Tracking page updates with the current state, technician, ETA/location information, and status history. Updates are persisted and available after re-login.'],
        ['9. Payment and completion', 'The client follows the selected payment route. The company records the payment where applicable and completes the job workflow.', 'Payment and booking status are saved. Company revenue/finance views and Super Admin finance oversight use the stored payment record.'],
        ['10. Client review and oversight', 'The client opens the completed booking, submits a rating/review, and can later see it in profile history. The company may reply. Super Admin can read all reviews and audit history.', 'The review links to the customer, company, and booking. Ratings/reviews appear in authorized client, company, and Admin views, while the Admin can take governance action if required.'],
    ], [1350, 4000, 4010])
    add_callout(doc, 'Visibility rule', 'A company becomes visible to clients only after Super Admin approval and while its company/account status remains active. When suspended, the same company is removed from client discovery and its company operations are blocked until reactivated.', 'FFF7E6')

    add_heading(doc, '6. Page Flow by Portal', 1)
    add_heading(doc, '6.1 Client Portal', 2)
    add_table(doc, ['Page / area', 'What the client does', 'Data used'], [
        ['Landing, Login, Register', 'Enters the marketplace or creates/signs into a client account.', 'User, session'],
        ['Browse Companies', 'Searches approved companies by text, city, area, or service filter; reads company descriptions and opens cards.', 'Company, Service, City, Review summary'],
        ['Company Details', 'Reads the selected company profile, active offers, technician details, reviews, and public inventory.', 'Company, Service, Technician, Review, Inventory'],
        ['Customize Booking', 'Selects service, quantity/optional inventory, location, date or request details.', 'Service, Inventory, booking draft'],
        ['Booking Summary / Payment', 'Confirms booking data and selects cash or card/online checkout flow.', 'Booking, Payment'],
        ['Bookings and Live Tracking', 'Views booking history, real-time work status, assigned technician, ETA, map/tracking, and job progression.', 'Booking, Technician, tracking updates'],
        ['Messages', 'Starts or continues a conversation with the chosen company.', 'ChatMessage, booking/conversation access'],
        ['Reviews and Profile', 'Edits own profile/photo, sees booking-linked reviews, and submits an eligible review.', 'User, Booking, Review'],
    ], [2100, 3860, 3400])
    add_heading(doc, '6.2 Company Portal', 2)
    add_table(doc, ['Page / area', 'What the company does', 'Data used'], [
        ['Overview', 'Sees own booking, revenue, job, and operational summary.', 'Company-scoped Booking, Payment, Technician, Inventory'],
        ['Company Details', 'Updates public name, location, areas, description, logo, and client-card cover image.', 'Company'],
        ['Requests / Bookings', 'Reviews incoming client requests and manages permitted status changes.', 'Booking'],
        ['Staff', 'Creates, edits, deletes technicians and uploads their images.', 'Technician'],
        ['Inventory', 'Creates, edits, deletes stock items used as optional booking items.', 'Inventory'],
        ['Services', 'Creates, edits, activates/deactivates, and deletes company offers.', 'Service'],
        ['Customers', 'Manages own customer records only.', 'Customer'],
        ['Chat, Payments, Reviews, Analytics', 'Communicates with clients, records / views payments, replies to reviews, and reviews company metrics.', 'ChatMessage, Payment, Review, Booking'],
    ], [2100, 3860, 3400])
    add_heading(doc, '6.3 Super Admin Portal', 2)
    add_table(doc, ['Page / area', 'What Super Admin does', 'Data used'], [
        ['Overview', 'Monitors platform totals and recent activity.', 'Aggregated platform records'],
        ['Registration Requests and Companies', 'Approves/rejects requests, inspects documents, blocks/reactivates companies.', 'Company, User, AuditEvent'],
        ['Ratings & Reviews', 'Reads platform-wide customer feedback and company responses.', 'Review, Company, User'],
        ['Users, Jobs, Finance, Support', 'Oversees user status, all bookings, payment records, and support requests.', 'User, Booking, Payment, SupportRequest'],
        ['Audit and Security Settings', 'Reviews sensitive actions and changes Super Admin credentials securely.', 'AuditEvent, User'],
    ], [2100, 4100, 3160])

    add_heading(doc, '7. Booking, Tracking, Payment, Review, and Chat Rules', 1)
    add_table(doc, ['Capability', 'Rule'], [
        ['Booking lifecycle', 'Pending -> Assigned -> En Route -> Arrived -> In Progress -> Completed -> Paid. The interface only presents actions valid for the current state.'],
        ['Cancellation', 'Clients may cancel only Pending or Assigned jobs. Company cancellation requires a reason.'],
        ['Technician assignment', 'The company may assign only a technician belonging to its own company and available for work.'],
        ['Tracking', 'Only the client and the related company can read booking tracking. Only the related approved company can update it.'],
        ['Payment', 'The client can select cash or request card/online checkout. The company can record payment against its own booking. Prices are resolved from the selected active service stored in MongoDB.'],
        ['Review', 'A customer creates a review against an eligible completed booking. The related company can reply; the Admin sees all reviews.'],
        ['Chat', 'Client and company access is restricted to the selected company conversation or booking relationship. Messages are saved to MongoDB and delivered through Socket.IO when connected.'],
    ], [2300, 7060])

    add_heading(doc, '8. Local Setup and Start Procedure', 1)
    add_numbered(doc, [
        'Install Node.js 20+ and MongoDB Community Server. MongoDB Compass alone is only a viewer; it does not start MongoDB.',
        'Start MongoDB Community Server and verify that port 27017 is available.',
        'Open PowerShell in the FleetOS project folder.',
        'Run npm install once. This installs dependencies for the root, server, client, and Admin applications.',
        'Run npm run dev. Wait until the API reports MongoDB connected and the client/Admin Vite addresses are displayed.',
        'Open client/company portal at the client address printed by Vite (normally http://localhost:5173) and Admin at the Admin address (normally http://localhost:5174).',
    ])
    add_callout(doc, 'Port conflict', 'If port 5000, 5173, or 5174 is already in use, stop the older FleetOS process first. Vite may move the UI to another port, but the API must be available before the portal can load real data.', 'FFF7E6')

    add_heading(doc, '9. Full Manual Test Plan', 1)
    add_text(doc, 'Before testing, use a clean browser session or separate browser profiles for Client, Company, and Super Admin. Keep MongoDB running throughout the test. Record the IDs, company names, and booking reference numbers created during the test.', after=6)
    add_heading(doc, '9.1 Pre-test checklist', 2)
    add_bullets(doc, [
        'MongoDB server is running and the API log says MongoDB connected.',
        'Client and Admin portals both load without a Vite error overlay.',
        'Use a new client email and a new company email for the test so results are easy to identify.',
        'Keep an approved test company with at least one active service and one available technician, or create these during the company test.',
        'Use valid image files under the configured image size limit for logo, cover, client profile, and technician photo checks.',
    ])
    add_heading(doc, '9.2 Company approval and public listing test', 2)
    add_numbered(doc, [
        'Register a new company with logo, business license, Pakistan city, service areas, business description, and credentials. Confirm the UI reports successful submission.',
        'Attempt to open company operations before approval. Expected: operational pages/API actions are blocked with approval-required messaging.',
        'Sign in to Super Admin, open Registration Requests, review the company, and approve it with a reason.',
        'Sign in to the newly approved company. Update Company Details: description, logo, client-card cover, city, areas, and location. Save and refresh the page.',
        'Add at least two active services in different categories, one technician with photo, and one inventory item linked to a relevant service category.',
        'Sign in as a client. Search by the company description, city, area, and each service category. Expected: the public card displays the saved company description, own logo/cover, actual services, and no unrelated car/vehicle wording.',
        'In Super Admin, suspend the company with a reason. Expected: company loses operations access and disappears from client discovery. Reactivate it and confirm it returns.',
    ])
    add_heading(doc, '9.3 Client-to-company booking test', 2)
    add_numbered(doc, [
        'Create a client account and complete the profile, including a picture if required.',
        'Browse the approved company, open Details, select an active service, add optional inventory if applicable, and enter the booking information.',
        'Confirm the booking and choose Cash or the online checkout route. Save the booking reference.',
        'Open the client Messages page and send a message to the company. Refresh the client page and confirm the message remains.',
        'Open the company Requests/Bookings page. Expected: the same booking appears for that company only, with the client information and requested service.',
        'Accept/update the booking, assign the technician, then update the tracking/status one step at a time.',
        'Open the client live tracking page. Expected: booking status, assigned technician, ETA/tracking information, and company identity update correctly.',
        'Complete the job, record payment if applicable, and check that the payment appears in company payments and Super Admin finance.',
        'From the client completed booking, submit a rating and review. Expected: it appears in client Reviews, company Reviews, technician/company presentation where supported, and Super Admin Ratings & Reviews.',
    ])
    add_heading(doc, '9.4 CRUD and isolation test', 2)
    add_table(doc, ['Feature', 'Create', 'Read after refresh', 'Update', 'Delete / restriction'], [
        ['Company profile', 'Save logo, cover, description, areas', 'Client card and company form match stored values', 'Change description/logo and verify client refresh', 'Only company owner scope may update'],
        ['Services', 'Add active service', 'Appears in company and client booking options', 'Change price/category/status', 'Deleted/inactive service cannot be selected'],
        ['Technicians', 'Add technician + photo', 'Appears in staff and assignment list', 'Change availability/details', 'Removed technician cannot be assigned'],
        ['Inventory', 'Add stock item', 'Appears only for its company/service context', 'Change quantity/price', 'Deleted item no longer offered'],
        ['Bookings', 'Client creates', 'Both authorized sides see same reference', 'Company updates valid status only', 'Client cancellation restricted by status'],
        ['Messages', 'Send from client/company', 'Refresh both sides; message persists', 'New messages arrive in conversation', 'Unauthorized account cannot open conversation'],
        ['Reviews', 'Client submits after eligible job', 'All authorized views match', 'Company replies', 'Duplicate/unauthorized review is rejected'],
    ], [1620, 1800, 2400, 1800, 1740])

    add_heading(doc, '10. MongoDB Persistence Verification', 1)
    add_text(doc, 'Use MongoDB Compass or mongosh only as a verification tool after performing UI actions. Do not edit production records directly from Compass during normal operation.', after=6)
    add_table(doc, ['UI action', 'Collection to inspect', 'Expected evidence'], [
        ['Client or company registration', 'users, companies', 'New email, role, hashed password (not plain password), company link, and pending/active state.'],
        ['Company public profile save', 'companies', 'Updated description, logo data/path, hero image, city, areas, location, and updatedAt.'],
        ['Service / staff / inventory change', 'services, technicians, inventories', 'Company ObjectId matches signed-in company; changed record and timestamps are present.'],
        ['Client booking', 'bookings', 'Booking reference, customer/company/service references, service snapshot, status history, totals, and payment method.'],
        ['Message sent', 'chatmessages', 'Sender, conversation/booking reference, text, and timestamp persist after browser refresh.'],
        ['Payment and review', 'payments, reviews', 'Correct booking/company/customer references and recorded amounts/rating/comment.'],
        ['Admin approval/suspension', 'companies, users, auditevents', 'New approval/status plus audit reason, actor, timestamp, and action.'],
    ], [2000, 2200, 5160])
    add_callout(doc, 'Data integrity check', 'For any created record, compare the MongoDB company/customer/booking references with the ID shown in the API response or UI flow. A record belonging to a different company is a critical authorization defect.', 'FFF0F0')

    add_heading(doc, '11. Automated Checks and Build Verification', 1)
    add_table(doc, ['Command', 'Purpose', 'Expected result'], [
        ['npm run build', 'Build client and Super Admin production bundles.', 'Both Vite builds complete without errors.'],
        ['npm test', 'Run server contract/schema tests.', 'All Node test cases pass.'],
        ['npm --prefix server run audit:mongo', 'Exercise persistence audit when MongoDB is available.', 'CRUD and key MongoDB workflows pass.'],
        ['npm run dev', 'Run API, client/company portal, and Admin portal together.', 'API becomes healthy; portal URLs are displayed.'],
    ], [3000, 3500, 2860])
    add_text(doc, 'Recommended regression sequence', 12, NAVY, True, before=8, after=4)
    add_bullets(doc, [
        'Run the production build before every delivery.',
        'Run server tests before every backend change.',
        'Run the MongoDB persistence audit with MongoDB active after changes to controllers, models, auth, bookings, payments, reviews, or chat.',
        'Perform the three-browser manual flow in Section 9 after changes to UI routing, roles, approval, or real-time features.',
    ])

    add_heading(doc, '12. Deployment Readiness', 1)
    add_bullets(doc, [
        'Use a managed MongoDB deployment such as MongoDB Atlas, and set its full URI as MONGODB_URI in the API host environment.',
        'Set NODE_ENV=production, a strong JWT_SECRET, and CORS_ORIGINS containing the exact client and Admin deployment URLs.',
        'Deploy the Express + Socket.IO server where WebSocket connections are supported and expose the API over HTTPS.',
        'Deploy the client/company and Super Admin builds as separate static web applications, each configured to call the deployed API.',
        'Do not store Super Admin credentials or plain passwords in source code. Use the approved credential setup/rotation path so MongoDB contains only bcrypt password hashes.',
        'Confirm production startup logs show a successful database connection. Production mode deliberately avoids modifying development demonstration records.',
        'Configure database backups, least-privilege database credentials, log monitoring, and an uptime check for the API health endpoint.',
    ])
    add_heading(doc, '13. Common Troubleshooting', 1)
    add_table(doc, ['Symptom', 'Likely cause', 'Resolution'], [
        ['ECONNREFUSED 127.0.0.1:27017', 'MongoDB server is stopped.', 'Start MongoDB Community Server, then restart npm run dev. Compass alone is not sufficient.'],
        ['EADDRINUSE on 5000 / 5173 / 5174', 'An old API or Vite process is still active.', 'Stop the old process, then start FleetOS again. Use the printed Vite address if a UI port changes.'],
        ['Company cannot open operational pages', 'Company is pending, rejected, suspended, or not the company role.', 'Check Super Admin company status and sign in again after approval/reactivation.'],
        ['Company absent from client listing', 'Not approved/active, no matching filter, or stale browser data.', 'Check Admin status, clear filters, refresh, and verify the company record/service in MongoDB.'],
        ['Chat or tracking missing', 'User is not authorized for that booking/conversation or API/socket is offline.', 'Confirm both accounts relate to the same booking/company and API is running.'],
        ['Text/logo display issue', 'Cached frontend bundle or an older dev process.', 'Hard refresh the browser and restart the affected Vite portal after confirming the latest build.'],
    ], [2250, 2600, 4510])

    add_heading(doc, '14. Final Acceptance Checklist', 1)
    add_bullets(doc, [
        'Client, company, and Super Admin can sign in only through their intended role paths.',
        'A new company follows pending -> approved -> client-visible and suspension removes visibility/access.',
        'All company profile, services, technicians, inventory, customer, booking, payment, message, and review CRUD changes persist after refresh and re-login.',
        'Client discovery uses the company’s own description, branding, location, active services, and Pakistan-only city/area data.',
        'Booking, assignment, tracking, payment, review, and chat work across the correct client/company pair only.',
        'Super Admin sees platform-wide records and an audit entry for sensitive changes.',
        'Build and server tests pass; MongoDB persistence audit passes with the target database active.',
        'Production configuration uses managed MongoDB, HTTPS, secrets in environment variables, CORS origins, and backups.',
    ])
    add_callout(doc, 'Sign-off recommendation', 'Use the Section 9 manual test plan with a real new company and client account, then attach the completed test evidence (screenshots, booking reference, and MongoDB verification) to the project submission.', 'EAF0F7')

def main():
    doc = Document()
    setup(doc)
    cover(doc)
    document_body(doc)
    doc.core_properties.title = 'FleetOS System Documentation and Test Guide'
    doc.core_properties.subject = 'FleetOS architecture, permissions, workflow, deployment, and testing'
    doc.core_properties.author = 'FleetOS Project Team'
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    main()
